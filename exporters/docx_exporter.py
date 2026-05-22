"""
Generate a Word document version of the tech pack using python-docx.
"""

import base64
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def _add_kv_table(doc: Document, items: list):
    """Add a 2-column key-value table."""
    table = doc.add_table(rows=len(items), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)
    for i, (label, value) in enumerate(items):
        row = table.rows[i]
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)
        c0 = row.cells[0].paragraphs[0]
        run = c0.add_run(str(label))
        run.bold = True
        run.font.size = Pt(10)
        c1 = row.cells[1].paragraphs[0]
        run = c1.add_run(str(value) if value is not None else "—")
        run.font.size = Pt(10)


def _add_measurement_table(doc: Document, measurements: dict, base_size: str):
    table = doc.add_table(rows=1 + len(measurements), cols=3)
    table.style = "Light Grid"
    header = table.rows[0].cells
    header[0].text = "Measurement Point"
    header[1].text = f"Size {base_size or 'M'}"
    header[2].text = "Tolerance"
    for cell in header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for i, (point, vals) in enumerate(measurements.items(), start=1):
        row = table.rows[i].cells
        row[0].text = point
        row[1].text = f"{vals.get('value', 0)} cm"
        row[2].text = f"± {vals.get('tolerance', 0)} cm"
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def _sizes_for_range(size_range: str) -> list[str]:
    """Expand a size_range label into the column headers used in spec sheets."""
    return {
        "S - XL":     ["S", "M", "L", "XL"],
        "XS - XXL":   ["XS", "S", "M", "L", "XL", "XXL"],
        "XXS - XXXL": ["XXS", "XS", "S", "M", "L", "XL", "XXL", "XXXL"],
    }.get(size_range, ["S", "M", "L", "XL", "XXL"])


def _add_team_spec_sheet(doc: Document, data: dict):
    """Render a spec sheet section matching the team's RWS .xls template.

    Layout:
      - Header info (Buyer's style #, Composition, Gauge & Ends, Maker, etc.)
      - Multi-size measurement table: rows = measurement points,
        columns = sizes, base_size column populated, others blank.
    """
    # Import the measurement-point list here so we don't need a top-level import
    # cycle (this module's caller already has it via the form).
    from config.dropdown_options import KNITWEAR_MEASUREMENT_POINTS

    _add_h2(doc, "Spec Sheet (factory-ready)")

    # 9-row key-value header — same labels and order the team uses
    header_rows = [
        ("Buyer's style # :",  data.get("style_number") or ""),
        ("Factory style # :",  ""),
        ("Composition :",      data.get("composition") or ""),
        ("Final Content:",     ""),
        ("Gauge & Ends :",     f"{data.get('gauge') or ''}  {data.get('ends') or '2 ends'}".strip()),
        ("Remarks :",          "size spec"),
        ("Date :",             datetime.now().strftime("%Y-%b-%d")),
        ("Maker :",            data.get("maker") or "New world"),
        ("Unit:",              "CM"),
    ]
    info_table = doc.add_table(rows=len(header_rows), cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Cm(4.5)
    info_table.columns[1].width = Cm(10)
    for i, (label, value) in enumerate(header_rows):
        row = info_table.rows[i]
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(10)
        c0 = row.cells[0].paragraphs[0]
        run = c0.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        c1 = row.cells[1].paragraphs[0]
        run = c1.add_run(str(value))
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Multi-size measurement table
    sizes = _sizes_for_range(data.get("size_range") or "S - XL")
    base_size = data.get("base_size") or "M"
    measurements = data.get("measurements") or {}

    n_cols = 3 + len(sizes)  # # | Measurement Point | Description | <sizes>
    n_rows = 1 + len(KNITWEAR_MEASUREMENT_POINTS)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid"

    # Header row
    header = table.rows[0].cells
    header[0].text = "#"
    header[1].text = "Measurement Point"
    header[2].text = "Description"
    for i, sz in enumerate(sizes):
        header[3 + i].text = sz
    for cell in header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # One row per known measurement point
    base_col = sizes.index(base_size) if base_size in sizes else len(sizes) // 2
    for i, (point, _unit, _tol) in enumerate(KNITWEAR_MEASUREMENT_POINTS, start=1):
        row = table.rows[i].cells
        row[0].text = f"{i}.0"
        # Split "Front body Length (fm HPS)" into name + description in parens
        if "(" in point and point.endswith(")"):
            name, desc = point.rsplit("(", 1)
            row[1].text = name.strip()
            row[2].text = desc.rstrip(")").strip()
        else:
            row[1].text = point
            row[2].text = ""
        # Fill base_size column if we have a value
        val = measurements.get(point, {}).get("value") if isinstance(measurements.get(point), dict) else None
        if val:
            row[3 + base_col].text = f"{val}"
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _add_image_grid(doc: Document, images: list, per_row: int = 2):
    """Insert the images as a 2-column table of (image / caption) cells."""
    cell_width_cm = 7.5

    # Pad to a multiple of `per_row` so the last row stays even.
    padded = list(images) + [None] * ((per_row - len(images) % per_row) % per_row)
    row_count = len(padded) // per_row
    if row_count == 0:
        return

    table = doc.add_table(rows=row_count * 2, cols=per_row)  # 2 rows per image: image, caption
    table.autofit = False

    for ri in range(row_count):
        img_row = table.rows[ri * 2].cells
        cap_row = table.rows[ri * 2 + 1].cells
        for ci in range(per_row):
            entry = padded[ri * per_row + ci]
            img_cell = img_row[ci]
            cap_cell = cap_row[ci]
            img_cell.width = Cm(cell_width_cm)
            cap_cell.width = Cm(cell_width_cm)
            if entry is None:
                continue
            # Insert the image
            try:
                raw = base64.b64decode(entry["data"])
            except Exception:
                continue
            p = img_cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(BytesIO(raw), width=Cm(cell_width_cm - 0.2))
            # Caption
            cap_p = cap_cell.paragraphs[0]
            cap_run = cap_p.add_run(entry.get("caption") or "")
            cap_run.italic = True
            cap_run.font.size = Pt(9)

    # Space after the grid
    doc.add_paragraph()


def generate_docx(data: dict) -> bytes:
    """Render the tech pack as a .docx file and return raw bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TECH PACK")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"{data.get('style_name') or '—'} · {data.get('style_number') or '—'}"
        f" · {data.get('season') or '—'}"
    )
    sub_run.font.size = Pt(11)

    timestamp = doc.add_paragraph()
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = timestamp.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ts_run.font.size = Pt(9)
    ts_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    is_knitwear = (data.get("product_type") or "").startswith("Knitwear")

    # Spec sheet (team factory-ready format) — produced first so factories
    # who only need the standard layout can stop reading right there.
    if is_knitwear:
        _add_team_spec_sheet(doc, data)

    # Images (placed before the textual sections so they're seen first)
    images = data.get("images") or []
    if images:
        _add_h2(doc, "Images & References")
        _add_image_grid(doc, images)

    # AI Technical Drawing — its own section
    tech_drawing = data.get("technical_drawing")
    if tech_drawing:
        _add_h2(doc, "Technical Drawing (AI-generated)")
        _add_image_grid(doc, [tech_drawing], per_row=1)

    # 1. Style Overview
    _add_h2(doc, "1. Style Overview")
    _add_kv_table(doc, [
        ("Product Type", data.get("product_type")),
        ("Style Name", data.get("style_name")),
        ("Style Number", data.get("style_number")),
        ("Season", data.get("season")),
        ("Gender", data.get("gender")),
        ("Fit", data.get("fit")),
        ("Size Range", data.get("size_range")),
        ("Color", data.get("color_name")),
        ("Pantone Code", data.get("pantone_code")),
        ("Composition", data.get("composition")),
    ])

    # 2. Material & Construction
    _add_h2(doc, "2. Material & Construction")
    if is_knitwear:
        _add_kv_table(doc, [
            ("Yarn Type", data.get("yarn_type")),
            ("Yarn Count", data.get("yarn_count")),
            ("Gauge", data.get("gauge")),
            ("Knit Structure", data.get("knit_structure")),
            ("Rib Structure", data.get("rib_structure")),
            ("Dyeing Method", data.get("dye_method")),
        ])
    else:
        _add_kv_table(doc, [
            ("Fabric Structure", data.get("fabric_structure")),
            ("Fabric Weight (gsm)", data.get("fabric_weight_gsm")),
            ("Dyeing Method", data.get("dye_method")),
        ])

    # 3. Style Details
    _add_h2(doc, "3. Style Details")
    details = [
        ("Neckline", f"{data.get('neckline') or '—'} (rib: {data.get('neckline_rib_cm') or 0} cm)"),
        ("Sleeve", f"{data.get('sleeve_length') or '—'} / {data.get('sleeve_type') or '—'}"),
        ("Hem", f"{data.get('hem_style') or '—'} ({data.get('hem_height_cm') or 0} cm)"),
        ("Cuff", f"{data.get('cuff_style') or '—'} ({data.get('cuff_height_cm') or 0} cm)"),
        ("Placket / Closure", data.get("placket")),
        ("Placket interlining", data.get("placket_interlining")),
    ]
    if "button" in (data.get("placket") or "").lower():
        details.append((
            "Buttons",
            f"{data.get('button_count') or 0} × {data.get('button_size_l') or '—'} "
            f"{data.get('button_material') or ''} ({data.get('button_color') or ''})",
        ))
    details.extend([
        ("Print / Embroidery", data.get("print_embroidery")),
        ("Wash / Finishing", data.get("wash_finishing")),
        ("Stitching / Construction", data.get("stitching_type")),
        ("Shoulder Reinforcement", "Yes" if data.get("shoulder_reinforcement") else "No"),
    ])
    _add_kv_table(doc, details)

    # 4. Measurements
    _add_h2(doc, f"4. Measurements (Size {data.get('base_size') or 'M'})")
    measurements = data.get("measurements") or {}
    if measurements:
        _add_measurement_table(doc, measurements, data.get("base_size") or "M")
    else:
        doc.add_paragraph("No measurements specified.")

    # 5. Labels & Packing
    _add_h2(doc, "5. Labels & Packing")
    _add_kv_table(doc, [
        ("Labels", ", ".join(data.get("labels") or []) or "—"),
        ("Packing", data.get("packing")),
    ])

    # 6. Supplier Actions
    _add_h2(doc, "6. Supplier Actions Required")
    actions = data.get("supplier_actions") or []
    if actions:
        for a in actions:
            doc.add_paragraph(a, style="List Bullet")
    else:
        doc.add_paragraph("None specified.")

    # 7. Commercial
    _add_h2(doc, "7. Commercial Information")
    _add_kv_table(doc, [
        ("Target Quantity", f"{data.get('target_quantity') or 0} pcs"),
        ("Target Price", f"${data.get('target_price_usd') or 0:.2f} USD"),
        ("Delivery Date", str(data.get("delivery_date") or "—")),
    ])
    if data.get("notes"):
        notes_p = doc.add_paragraph()
        run = notes_p.add_run("Notes: ")
        run.bold = True
        notes_p.add_run(str(data["notes"]))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
